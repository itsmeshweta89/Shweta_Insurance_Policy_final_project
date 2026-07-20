"""One-off generator for the synthetic InsureWise Policy Co-Pilot knowledge base.

Run once: `python -m data.generate_sample_docs`. Output is committed to the
repo under data/knowledge_base/ so docker-compose / a fresh checkout doesn't
need to regenerate it before running `python -m app.rag.ingest`.
"""
import csv
import os

from docx import Document
from fpdf import FPDF

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE_DIR, "knowledge_base")


def write_life_insurance_brochure():
    html = """<!DOCTYPE html>
<html>
<head><title>InsureWise SecureLife Term Plan - Product Brochure</title></head>
<body>
<h1>InsureWise SecureLife Term Plan</h1>

<h2>Plan Overview</h2>
<p>SecureLife is a pure term life insurance plan offering a high sum
assured at an affordable premium. It is available to individuals aged
18 to 65 years, with cover continuing up to age 75. The minimum sum
assured is Rs. 25 lakh and there is no upper limit, subject to
underwriting.</p>

<h2>Premium Payment Options</h2>
<p>Policyholders can choose Regular Pay (annual, half-yearly, or
monthly via ECS/NACH), Limited Pay (premiums payable for 10 or 15
years while cover continues for the full policy term), or Single Pay.
A 2% discount on the annualized premium is offered for annual mode
compared to monthly mode. Female lives and non-tobacco users receive
preferential premium rates.</p>

<h2>Riders Available</h2>
<p>SecureLife can be enhanced with the Accidental Death Benefit Rider,
Critical Illness Rider (covering 20 listed illnesses), and Waiver of
Premium Rider. Total rider premium cannot exceed 30% of the base plan
premium as per IRDAI (Insurance Regulatory and Development Authority
of India) guidelines.</p>

<h2>Death Benefit</h2>
<p>On death of the life assured during the policy term, the nominee
receives the highest of: (a) the sum assured, (b) 10 times the
annualized premium, or (c) 105% of total premiums paid till date of
death, provided all due premiums have been paid.</p>

<h2>Free-Look Period and Grace Period</h2>
<p>A free-look period of 30 days from the date of receipt of the
policy document is available for policies purchased through
distance-marketing channels, and 15 days for all other channels. A
grace period of 30 days is allowed for annual, half-yearly, and
quarterly premiums, and 15 days for monthly premiums.</p>

<h2>Exclusions</h2>
<p>Death due to suicide within 12 months of the policy commencement or
revival date is excluded, except that 80% of premiums paid (or the
surrender value, if higher) will be payable to the nominee in such
cases, as per IRDAI (Protection of Policyholders' Interests)
Regulations.</p>
</body>
</html>
"""
    path = os.path.join(OUT_DIR, "product_brochure_securelife_term.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print("wrote", path)


def write_health_insurance_brochure():
    html = """<!DOCTYPE html>
<html>
<head><title>InsureWise HealthShield Family Floater - Product Brochure</title></head>
<body>
<h1>InsureWise HealthShield Family Floater Plan</h1>

<h2>Sum Insured Options</h2>
<p>HealthShield is available on a family floater basis covering
self, spouse, and up to 3 dependent children, with sum insured
options of Rs. 3 lakh, 5 lakh, 10 lakh, 25 lakh, and 50 lakh. Entry
age is 18 to 65 years for adults and 91 days to 25 years for
dependent children.</p>

<h2>Waiting Periods</h2>
<p>An initial waiting period of 30 days applies to all illnesses
except accidental hospitalization. Pre-existing diseases are covered
after a waiting period of 36 months of continuous coverage. Specific
illnesses such as cataract, hernia, and joint replacement surgery
have a waiting period of 24 months.</p>

<h2>No Claim Bonus</h2>
<p>A No Claim Bonus (NCB) of 10% of the sum insured is added for each
claim-free year, up to a maximum of 50% of the base sum insured. In
the event of a claim, the NCB reduces by 10% at the next renewal but
the base sum insured remains unaffected.</p>

<h2>Room Rent and Sub-limits</h2>
<p>For sum insured up to Rs. 5 lakh, room rent is capped at 1% of the
sum insured per day, and ICU charges at 2% per day. No room-rent
capping applies for sum insured of Rs. 10 lakh and above. Cataract
surgery is capped at Rs. 40,000 per eye for sum insured up to Rs. 5
lakh.</p>

<h2>Cashless Network and Claims</h2>
<p>HealthShield offers cashless treatment at over 8,500 network
hospitals across India. Pre-authorization for planned hospitalization
should be sought at least 48 hours in advance; emergency
hospitalization must be intimated within 24 hours of admission.
Reimbursement claims must be filed within 30 days of discharge.</p>

<h2>Tax Benefit</h2>
<p>Premiums paid towards HealthShield qualify for deduction under
Section 80D of the Income Tax Act, 1961, subject to the limits
prescribed for individuals and senior citizens.</p>
</body>
</html>
"""
    path = os.path.join(OUT_DIR, "product_brochure_healthshield_floater.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print("wrote", path)


def write_csr_disclosure():
    doc = Document()
    doc.add_heading("InsureWise Life Insurance Co. Ltd.", level=1)
    doc.add_heading(
        "Corporate Social Responsibility (CSR) Disclosure - FY 2024-25 "
        "(as per IRDAI Corporate Governance Guidelines)",
        level=2,
    )

    doc.add_heading("CSR Policy and Committee", level=2)
    doc.add_paragraph(
        "InsureWise's CSR Committee comprises three Board members, "
        "including one Independent Director, in compliance with "
        "Section 135 of the Companies Act, 2013 and IRDAI's Corporate "
        "Governance Guidelines for Insurers. The Committee meets at "
        "least twice a year to review CSR strategy and monitor "
        "implementation."
    )

    doc.add_heading("CSR Obligation and Spend", level=2)
    doc.add_paragraph(
        "As required under Section 135, InsureWise is obligated to "
        "spend at least 2% of the average net profits of the preceding "
        "three financial years on CSR activities. For FY 2024-25, the "
        "prescribed CSR obligation was Rs. 18.4 crore, against which "
        "the company spent Rs. 18.6 crore, resulting in a surplus "
        "carry-forward of Rs. 0.2 crore to the next financial year."
    )

    doc.add_heading("Focus Areas", level=2)
    doc.add_paragraph(
        "1. Financial literacy and insurance awareness programs in "
        "underserved rural districts (32% of CSR spend)\n"
        "2. Healthcare access, including mobile health camps and "
        "support for cancer and cardiac care (28% of CSR spend)\n"
        "3. Girl-child education scholarships and skill development "
        "(21% of CSR spend)\n"
        "4. Disaster relief and rehabilitation support (11% of CSR "
        "spend)\n"
        "5. Environmental sustainability, including afforestation "
        "drives (8% of CSR spend)"
    )

    doc.add_heading("Implementation Partners", level=2)
    doc.add_paragraph(
        "CSR programs are implemented directly and through registered "
        "implementing agencies empanelled under Schedule VII of the "
        "Companies Act, 2013. All implementing partners are subject to "
        "annual impact assessment as mandated for CSR projects with "
        "outlay exceeding Rs. 1 crore."
    )

    doc.add_heading("Board Responsibility Statement", level=2)
    doc.add_paragraph(
        "The Board of Directors confirms that the CSR policy has been "
        "implemented in accordance with the CSR objectives and policy "
        "of the company, and that the implementation and monitoring of "
        "CSR activities is in compliance with CSR objectives and the "
        "CSR Policy approved by the Board."
    )

    path = os.path.join(OUT_DIR, "csr_disclosure_fy2024_25.docx")
    doc.save(path)
    print("wrote", path)


def write_tax_circular():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(
        0, 10, "Internal Tax Circular: Taxation of Life and Health "
        "Insurance Products"
    )
    pdf.ln(4)

    sections = [
        (
            "Section 80C - Life Insurance Premiums",
            "Premiums paid towards life insurance policies qualify for "
            "deduction under Section 80C of the Income Tax Act, 1961, "
            "up to an overall limit of Rs. 1,50,000 per annum, combined "
            "with other eligible investments. For policies issued on or "
            "after 1 April 2012, the deduction is available only if the "
            "annual premium does not exceed 10% of the sum assured (15% "
            "for policies on the life of a person with disability or "
            "specified disease).",
        ),
        (
            "Section 10(10D) - Maturity Proceeds",
            "Maturity or death benefit proceeds from a life insurance "
            "policy are exempt under Section 10(10D), provided the "
            "premium-to-sum-assured ratio condition under Section 80C "
            "is met. Effective FY 2023-24, maturity proceeds from "
            "life insurance policies (other than ULIPs) issued on or "
            "after 1 April 2023 are taxable if the aggregate annual "
            "premium across all such policies exceeds Rs. 5,00,000.",
        ),
        (
            "Section 80D - Health Insurance Premiums",
            "Health insurance premiums are deductible up to Rs. 25,000 "
            "per annum for self, spouse, and dependent children, and an "
            "additional Rs. 25,000 for parents below 60 years. The "
            "limit increases to Rs. 50,000 where the insured person is a "
            "senior citizen (60 years or above). Preventive health "
            "check-up expenses up to Rs. 5,000 are included within "
            "these overall limits.",
        ),
        (
            "TDS on Insurance Payouts",
            "Under Section 194DA, TDS at 5% is deducted on the taxable "
            "portion of life insurance payouts (i.e., where proceeds are "
            "not exempt under Section 10(10D)) exceeding Rs. 1,00,000 in "
            "aggregate during a financial year. TDS is deducted on the "
            "income component (payout minus premiums paid), not the "
            "gross payout.",
        ),
        (
            "GST on Premiums",
            "GST is levied on life and health insurance premiums at "
            "18%, applied to the risk (mortality/morbidity) component "
            "of the premium. For traditional endowment plans, GST is "
            "charged at 4.5% of the premium in the first year and 2.25% "
            "in subsequent years, as per the composite rate rules for "
            "life insurance.",
        ),
        (
            "Disclaimer",
            "This circular is for internal reference only and "
            "summarizes provisions as understood as of the date of "
            "issue. It does not constitute tax advice. Customers should "
            "be directed to consult a qualified tax professional or "
            "chartered accountant for advice specific to their "
            "situation.",
        ),
    ]

    for title, body in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 8, title)
        pdf.set_font("Helvetica", "", 11)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, body)
        pdf.ln(2)

    path = os.path.join(OUT_DIR, "tax_circular_life_health_insurance.pdf")
    pdf.output(path)
    print("wrote", path)


def write_rider_documentation():
    rows = [
        ["rider_name", "applicable_base_plans", "entry_age", "max_cover", "premium_cap", "key_terms"],
        [
            "Accidental Death Benefit Rider",
            "SecureLife Term Plan, WealthBuilder Endowment",
            "18-60 years",
            "Equal to base sum assured, max Rs. 2 crore",
            "Rider premium max 30% of base premium (IRDAI cap)",
            "Pays additional sum assured on death due to accident within 180 days of the accident",
        ],
        [
            "Critical Illness Rider",
            "SecureLife Term Plan, HealthShield Floater",
            "18-55 years",
            "Max Rs. 50 lakh, capped at base sum assured",
            "Rider premium max 30% of base premium (IRDAI cap)",
            "Covers 20 listed critical illnesses; 90-day waiting period; survival period of 30 days post-diagnosis required for payout",
        ],
        [
            "Waiver of Premium Rider",
            "SecureLife Term Plan, WealthBuilder Endowment",
            "18-55 years",
            "Not applicable (waives future premiums)",
            "Rider premium approx. 3-5% of base premium",
            "Future premiums waived on diagnosis of total permanent disability or specified critical illness of the policyholder",
        ],
        [
            "Hospital Cash Rider",
            "HealthShield Floater",
            "18-65 years",
            "Rs. 1,000 to Rs. 5,000 per day of hospitalization",
            "Fixed premium slabs based on daily cash benefit chosen",
            "Daily cash benefit payable from the 2nd day of hospitalization, up to 30 days per policy year",
        ],
        [
            "Return of Premium Rider",
            "SecureLife Term Plan",
            "18-60 years",
            "Not applicable (returns premiums, not a cover amount)",
            "Increases base premium by approx. 40-60%",
            "Returns 100% of total premiums paid (excluding taxes and rider premiums) on survival to maturity, if all premiums are paid",
        ],
        [
            "Maternity and Newborn Cover Rider",
            "HealthShield Floater",
            "18-45 years (proposer)",
            "Rs. 50,000 to Rs. 2,00,000 per delivery",
            "Additional premium based on sum insured chosen",
            "Waiting period of 24-48 months depending on sum insured; covers normal and caesarean delivery plus newborn cover for 90 days",
        ],
    ]
    path = os.path.join(OUT_DIR, "rider_documentation.csv")
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print("wrote", path)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    write_life_insurance_brochure()
    write_health_insurance_brochure()
    write_csr_disclosure()
    write_tax_circular()
    write_rider_documentation()


if __name__ == "__main__":
    main()